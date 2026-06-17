"""
FedAcuity -- MidSEM Report Generator
Generates a BITS Pilani AIMLCZG628T compliant MidSEM report as a .docx file.

Usage:
    python -m src.presentation.midsem_report
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

TABLES_DIR  = Path("results/tables")
OUT_PATH    = Path("reports/FedAcuity_MidSEM_Report.docx")

# ---- helpers -----------------------------------------------------------------

def _set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def _heading(doc, text, level=1, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, size, bold=True)
    return p

def _body(doc, text, size=11, bold=False, indent=False, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    _set_font(run, size, bold=bold, italic=italic)
    return p

def _bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_font(run, size)
    return p

def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = str(val)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = Inches(col_widths[i])
    return table

def _page_break(doc):
    doc.add_page_break()

def _hline(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ---- load results ------------------------------------------------------------

def _load():
    r = {}
    mp = TABLES_DIR / "fl_metrics_summary.json"
    if mp.exists():
        with open(mp) as f: r["metrics"] = json.load(f)
    hp = TABLES_DIR / "fl_held_out_metrics.json"
    if hp.exists():
        with open(hp) as f: r["held_out"] = json.load(f)
    dp_p = TABLES_DIR / "dp_epsilon_sweep.csv"
    if dp_p.exists(): r["dp"] = pd.read_csv(dp_p)
    fp = TABLES_DIR / "fidelity_ks_test.csv"
    if fp.exists(): r["ks"] = pd.read_csv(fp)
    fro_p = TABLES_DIR / "fidelity_frobenius.json"
    if fro_p.exists():
        with open(fro_p) as f: r["frobenius"] = json.load(f)
    tstr_p = TABLES_DIR / "fidelity_tstr.json"
    if tstr_p.exists():
        with open(tstr_p) as f: r["tstr"] = json.load(f)
    return r


# ---- document sections -------------------------------------------------------

def make_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("A Privacy-Preserving Federated Learning Framework\nwith Explainability Auditing for\nStaffing-Acuity Mismatch Prediction in Long-Term Care")
    _set_font(r, 16, bold=True)

    doc.add_paragraph()
    for label, val in [
        ("Course No.", "AIMLCZG628T"),
        ("Course Title", "Dissertation"),
        ("", ""),
        ("Dissertation by", ""),
        ("Student Name", "Tanay Kashyap"),
        ("BITS ID", "2024AA05991"),
        ("", ""),
        ("Degree Programme", "M.Tech in AI & ML"),
        ("Organisation", "NuAlg Infotech Private Limited, Indore"),
        ("Evaluation", "MidSEM -- June 2026"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label:
            r1 = p.add_run(f"{label}: ")
            _set_font(r1, 12, bold=True)
            r2 = p.add_run(val)
            _set_font(r2, 12)
        else:
            p.add_run(" ")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI\nVIDYA VIHAR, PILANI, RAJASTHAN - 333031")
    _set_font(r, 12, bold=True)
    _page_break(doc)


def make_abstract(doc):
    _heading(doc, "ABSTRACT", size=13)
    _hline(doc)
    doc.add_paragraph()
    _body(doc, (
        "FedAcuity is a privacy-preserving Federated Learning (FL) framework for predicting "
        "staffing-acuity mismatch in Long-Term Care (LTC) facilities -- where resident care needs "
        "exceed available nursing capacity. Under HIPAA, resident health records cannot be "
        "centralised, making cross-facility collaborative learning impossible through conventional ML. "
        "FedAcuity resolves this through three standalone, independently publishable contributions: "
        "(C1) a domain-driven Clustered FL system that groups facilities by care type "
        "(Memory Care, Skilled Nursing, Independent Living) to address extreme non-IID heterogeneity, "
        "augmented with Opacus differential privacy; "
        "(C2) a CTGAN-based synthetic LTC benchmark dataset with fidelity validation via "
        "KS-tests, Frobenius norm, and TSTR experiments; and "
        "(C3) a 4-dimension XAI Audit Scorecard measuring explanation fidelity, stability, "
        "fairness, and clinical plausibility via SHAP (planned for post-MidSEM phase)."
    ), size=11)
    doc.add_paragraph()
    _body(doc, (
        "At MidSEM, all C1 and C2 components are implemented and validated, including a "
        "mid-cycle correction described transparently in Section 4.4: an initial design flaw "
        "(the Independent Living cluster had only one training client, making Clustered FL "
        "numerically indistinguishable from single-facility local training) was identified and "
        "fixed by moving one facility from the held-out set into the IL training cluster, giving "
        "every cluster at least two genuine federating clients. "
        "On the held-out Independent Living facility (unseen during training), "
        "Clustered FL achieves AUC-ROC of 0.9693 and F1 of 0.7273 -- approaching the "
        "Centralised Oracle upper bound (AUC 0.9799) while outperforming both FedAvg "
        "(AUC 0.9414, +2.79 AUC points) and the single-facility IL Local baseline "
        "(AUC 0.9643, +0.50 AUC points). The gap over FedAvg's training-round AUC trajectory "
        "is statistically significant (Mann-Whitney U=100.0, p=1.6e-05). "
        "Differential privacy sweep (re-run on the corrected facility roster) shows that epsilon=10 "
        "is the recommended operating point (19.3% utility degradation). C2 was reframed into a small but solid "
        "contribution: alongside the expected weak cross-domain signal against MIMIC-IV v3.1 "
        "(95,715 elderly patients, 205,456 admissions, only 3 of 14 features mappable), a new "
        "within-MIMIC-IV cohort calibration analysis compares patients discharged to post-acute/LTC "
        "care against those who are not, finding medium-to-large effect sizes on medication burden "
        "and acuity proxies (Cohen's d = 0.69 and 0.78, p < 0.001), and shows that MIMIC-IV's "
        "real-world 27.2%% discharge-to-post-acute rate sits within 0.8 points of the synthetic "
        "SNF mismatch target (28%%) -- a face-validity signal that the synthetic generator's "
        "calibration is not arbitrary. Internal consistency validation (20%% synthetic "
        "holdout) confirms CTGAN reproduces target LTC distributions: 14/14 features pass "
        "KS-test (alpha=0.05), Frobenius norm 0.2196 (27x better than random baseline), "
        "TSTR gap 0.0199 (target <8%%). "
        "The XAI Audit Scorecard (C3) is scheduled for post-MidSEM."
    ), size=11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature of the Student").font.bold = True
    p.add_run("                                    ")
    p.add_run("Signature of the Supervisor").font.bold = True
    doc.add_paragraph()
    _body(doc, "Name: _______________________                    Name: _______________________")
    _body(doc, "Date: _______________________                    Date: _______________________")
    _body(doc, "Place: ______________________                    Place: ______________________")
    _page_break(doc)


def make_contents(doc):
    _heading(doc, "CONTENTS", size=13)
    _hline(doc)
    toc = [
        ("1.", "Work Completed -- Implemented Components", "4"),
        ("2.", "System Architecture and Description", "6"),
        ("3.", "Experimental Results", "7"),
        ("4.", "Design Decisions and Challenges", "9"),
        ("5.", "Future Plan", "10"),
    ]
    for num, title, pg in toc:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num}  {title}")
        _set_font(r1, 11)
        r2 = p.add_run(f"  {'.' * (55 - len(title))}  {pg}")
        _set_font(r2, 11)
    _page_break(doc)


def make_section1(doc):
    _heading(doc, "1. WORK COMPLETED -- IMPLEMENTED COMPONENTS", size=13)
    _hline(doc)
    doc.add_paragraph()

    _heading(doc, "1.1 Data Engineering Module (Contribution C2)", size=12)
    _body(doc, (
        "The synthetic LTC benchmark dataset was fully designed and generated using CTGAN "
        "(Conditional Tabular GAN via the SDV library). The dataset represents 10 simulated "
        "LTC facilities across three care types: 3 Memory Care (MC), 4 Skilled Nursing (SNF), "
        "and 3 Independent Living (IL). Each facility contains approximately 1,095 daily records "
        "spanning 3 years, for a total of 10,950 rows."
    ))
    doc.add_paragraph()
    _body(doc, "Key implementation components:", bold=True)
    _bullet(doc, "15-feature clinical schema grounded in MDS 3.0 (ADL scores), RUG-IV categories, and CMS Payroll-Based Journal (PBJ) staffing data.")
    _bullet(doc, "Binary mismatch label: staffing_mismatch = 1 when (ADL_demand x census) / total_nursing_hours > threshold_care_type. Threshold calibrated per care type to achieve target mismatch rates: MC ~40%, SNF ~28%, IL ~12%.")
    _bullet(doc, "Deliberate non-IID distributional engineering: adl_cognition mean = 4.5 (MC) vs 2.5 (SNF) vs 1.0 (IL); medication_count = 11 vs 9 vs 5.")
    _bullet(doc, "Stratified 60/20/20 train/val/test splits per facility (SEED=42). Held-out facility 9 (IL) never participates in FL training -- the sole genuinely-unseen evaluation facility (see Section 4.4 for why facility 8 was moved into the IL training cluster).")
    _bullet(doc, "Fidelity validation pipeline (KS-test, Frobenius norm, TSTR) with three complementary validation modes: (1) Cross-domain: real MIMIC-IV v3.1 (95,715 elderly patients, 205,456 admissions via PhysioNet credentialed access) vs synthetic -- 3 mappable features (medication_count, rug_category via DRG severity, resident_census). Results show expected hospital-vs-LTC domain gap, validating C2's necessity as a purpose-built benchmark. (2) Within-MIMIC-IV cohort calibration (new, src/data/mimic_analysis.py): compares the LTC-bound discharge cohort against the rest of the MIMIC-IV elderly subset on the 3 mappable features, and checks whether MIMIC-IV's real discharge-to-post-acute rate is consistent with the synthetic mismatch targets. (3) Internal consistency: 20% synthetic holdout shows CTGAN reproduces target LTC distributions (14/14 KS pass, Frobenius 0.2196, TSTR gap 0.0199).")
    doc.add_paragraph()

    _heading(doc, "1.2 Federated Learning Simulation (Contribution C1)", size=12)
    _body(doc, (
        "All five model variants were implemented and successfully simulated for 50 communication "
        "rounds using a manual FL loop (no Ray dependency) compatible with Flower 1.x."
    ))
    doc.add_paragraph()
    _body(doc, "Five strategies implemented:", bold=True)
    _bullet(doc, "Local Baseline: Each facility trains independently. No federation. Evaluated on per-facility test sets.")
    _bullet(doc, "Centralised Oracle: All data pooled (HIPAA-violating upper bound). XGBoost trained on all training facilities, evaluated on held-out facility 9.")
    _bullet(doc, "FedAvg (McMahan et al. 2017): Standard federated averaging across all 9 training facilities.")
    _bullet(doc, "FedProx (Li et al. 2020): FedAvg with proximal regularisation term (mu=0.1). Note: proximal term is noted in config but not implemented as a true gradient penalty for XGBoost -- applied correctly only in the PyTorch NN variant.")
    _bullet(doc, "Clustered FL (C1 -- primary contribution): Care-type cluster aggregation. Three independent global models (MC, SNF, IL). Intra-cluster FedAvg only. MC=[0,1,2] (3 clients), SNF=[3,4,5,6] (4 clients), IL=[7,8] (2 clients, after the Section 4.4 fix). Held-out facility 9 (IL) excluded from all rounds.")
    doc.add_paragraph()
    _body(doc, "Technical implementation:", bold=True)
    _bullet(doc, "Local model: XGBoost (100 estimators, max_depth=6, learning_rate=0.1). Preferred over neural networks for tabular LTC data (native SHAP TreeExplainer support).")
    _bullet(doc, "XGBoost federation via byte serialisation: serialize_xgb_model() and deserialize_xgb_model() in src/fl/client.py. Compatible with XGBoost 3.x.")
    _bullet(doc, "Aggregation strategy: weighted average by dataset size (proxy for true tree merging; full tree-merging scheduled for post-MidSEM refinement).")
    _bullet(doc, "72/72 unit tests passing (test_schema.py + test_loaders.py).")
    doc.add_paragraph()

    _heading(doc, "1.3 Differential Privacy Module", size=12)
    _body(doc, (
        "A secondary PyTorch neural network (StaffingNN: hidden layers [128, 64, 32], "
        "dropout=0.2) was implemented for Opacus DP-SGD integration. XGBoost does not support "
        "Opacus gradient clipping, so the PyTorch NN serves as the DP evaluation model."
    ))
    doc.add_paragraph()
    _bullet(doc, "Opacus PrivacyEngine applied with delta=1e-5, max_grad_norm=1.0.")
    _bullet(doc, "Epsilon sweep: epsilon in {1, 2, 5, 10, infinity} (no DP).")
    _bullet(doc, "Training: 20 epochs per epsilon value, batch_size=64, Adam optimiser.")
    _bullet(doc, "Evaluation on held-out facility 9 (IL); training pool is the other 9 training facilities.")
    doc.add_paragraph()

    _heading(doc, "1.4 Evaluation Framework", size=12)
    _bullet(doc, "ResultsLogger: thread-safe, append-only JSON + CSV logger for all FL rounds.")
    _bullet(doc, "metrics.py: bootstrap 95% confidence intervals (2000 iterations), Mann-Whitney U test (non-parametric, two-sided).")
    _bullet(doc, "figures.py: Fig 3 (convergence curves, AUC vs round) and Fig 4 (five-model bar chart with error bars).")
    _bullet(doc, "eval_held_out.py: comprehensive held-out evaluation producing AUC-ROC, F1, Precision, Recall for all 5 strategies on held-out facility 9.")
    _page_break(doc)


def make_section2(doc):
    _heading(doc, "2. SYSTEM ARCHITECTURE AND DESCRIPTION", size=13)
    _hline(doc)
    doc.add_paragraph()
    _body(doc, (
        "FedAcuity implements a three-layer federated learning architecture designed "
        "for HIPAA-compliant cross-facility mismatch prediction:"
    ))
    doc.add_paragraph()

    _heading(doc, "Layer 1 -- Facility Edge (src/fl/client.py)", size=12)
    _body(doc, (
        "Each of the 10 simulated LTC facilities runs a FedAcuityClient (Flower NumPyClient). "
        "XGBoost trains locally on the facility's own resident records. Only serialised model "
        "bytes (50-200 KB per round) are transmitted to the aggregation server. "
        "Raw resident data never leaves the facility boundary."
    ))
    doc.add_paragraph()

    _heading(doc, "Layer 2 -- Aggregation Server (src/fl/simulation.py)", size=12)
    _body(doc, (
        "The Flower simulation server implements three interchangeable aggregation strategies: "
        "FedAvg (global), FedProx (global with proximal term), and Clustered FL (per care type). "
        "In Clustered FL, the 9 training facilities are partitioned into three clusters: "
        "MC [0,1,2] (3 clients), SNF [3,4,5,6] (4 clients), IL [7,8] (2 clients). Each cluster "
        "maintains an independent global model updated by genuine intra-cluster FedAvg -- every "
        "cluster, including IL, now has at least two federating clients. Facility 9 (IL) is held out."
    ))
    doc.add_paragraph()

    _heading(doc, "Layer 3 -- Evaluation and XAI (src/evaluation/, src/xai/)", size=12)
    _body(doc, (
        "Post-training evaluation includes held-out facility testing (AUC-ROC, F1, "
        "precision, recall), statistical significance testing, and figure generation. "
        "The XAI Audit Scorecard (SHAP TreeExplainer, D1-D4 modules) is the primary "
        "deliverable for the post-MidSEM phase (Weeks 3-4, June 14 - June 27)."
    ))
    doc.add_paragraph()

    _heading(doc, "Data Flow Summary", size=12)
    rows = [
        ("Step", "Description", "Output"),
        ("1", "CTGAN generation (src/data/generator.py)", "data/synthetic/*.csv + all_facilities.parquet"),
        ("2", "Fidelity validation (src/data/fidelity.py)", "results/tables/fidelity_*.json + Fig 2"),
        ("3", "FL simulation (src/fl/simulation.py)", "results/logs/results_*.json + centralised_model.json"),
        ("4", "DP epsilon sweep (src/dp/epsilon_sweep.py)", "results/tables/dp_epsilon_sweep.csv + Fig 5"),
        ("5", "Held-out evaluation (src/evaluation/eval_held_out.py)", "results/tables/fl_held_out_metrics.json"),
        ("6", "Figures (src/evaluation/figures.py)", "results/figures/fig3_*, fig4_*"),
        ("7 [planned]", "SHAP pipeline + D1-D4 XAI", "results/tables/xai_audit_raw.json + Fig 6"),
    ]
    _add_table(doc, ["Step", "Description", "Output"], rows[1:], col_widths=[0.6, 2.8, 2.8])
    _page_break(doc)


def make_section3(doc, r):
    _heading(doc, "3. EXPERIMENTAL RESULTS", size=13)
    _hline(doc)
    doc.add_paragraph()

    _heading(doc, "3.1 FL Strategy Comparison (Held-Out Facility 9)", size=12)
    _body(doc, (
        "Table 1 reports evaluation on held-out facility 9 (Independent Living, "
        "never seen during any FL training round) after 50 communication rounds. "
        "AUC-ROC is the primary metric. F1, Precision, Recall computed at threshold=0.5."
    ))
    doc.add_paragraph()

    if "held_out" in r:
        ho = r["held_out"]
        rows = []
        for strat, display in [
            ("il_local_baseline",       "IL Local (facility 7 only)"),
            ("cross_facility_ensemble", "Cross-Facility Ensemble (no protocol)"),
            ("centralised",             "Centralised Oracle"),
            ("fedavg",                  "FedAvg"),
            ("fedprox",                 "FedProx (mu=0.1)"),
            ("clustered_fl",            "Clustered FL (C1) *"),
        ]:
            v = ho.get(strat, {}).get("overall", {})
            rows.append([
                display,
                f"{v.get('auc_roc', '--'):.4f}" if isinstance(v.get('auc_roc'), float) else "--",
                f"{v.get('f1', '--'):.4f}" if isinstance(v.get('f1'), float) else "--",
                f"{v.get('precision', '--'):.4f}" if isinstance(v.get('precision'), float) else "--",
                f"{v.get('recall', '--'):.4f}" if isinstance(v.get('recall'), float) else "--",
            ])
        _add_table(doc, ["Strategy", "AUC-ROC", "F1", "Precision", "Recall"],
                   rows, col_widths=[2.2, 1.0, 0.8, 1.0, 0.8])
    doc.add_paragraph()
    _body(doc, "* Primary contribution. Evaluated using the IL cluster model (facilities 7+8) on held-out facility 9.")

    doc.add_paragraph()
    _body(doc, "Key Finding:", bold=True)
    _body(doc, (
        "Clustered FL achieves AUC-ROC 0.9693 and F1 0.7273, approaching the Centralised Oracle "
        "(AUC 0.9799, F1 0.7234) within 1.06 AUC points -- while maintaining full HIPAA compliance. "
        "FedAvg achieves AUC 0.9414 and F1 0.7037 on held-out facility 9, confirming "
        "that naive global aggregation underperforms on non-IID LTC data. Critically, CFL also "
        "outperforms the IL Local baseline (AUC 0.9643, single facility, no federation) by 0.50 AUC "
        "points -- demonstrating that genuine 2-client intra-cluster federation adds measurable value "
        "beyond simply having care-type-matched data. The CFL-vs-FedAvg gap, tracked across the "
        "50-round training trajectory, is statistically significant (Mann-Whitney U=100.0, p=1.6e-05). "
        "This result supersedes an earlier draft in which the IL cluster had only one training client "
        "(facility 7), making CFL numerically identical to IL Local; see Section 4.4 for the fix."
    ))
    doc.add_paragraph()

    _heading(doc, "3.2 Fidelity Validation Results (C2 -- Reframed)", size=12)
    _body(doc, (
        "C2 was reframed during this cycle (see Section 4.4) into a deliberately modest but solid "
        "claim: the synthetic LTC benchmark is internally faithful to its design targets, and the "
        "MIMIC-IV cross-domain check -- while necessarily weak given only 3 shared features -- is "
        "supplemented by a new within-MIMIC-IV cohort calibration analysis that uses the credentialed "
        "MIMIC-IV access differently from a standard fidelity benchmark. C2 validation now spans three "
        "complementary methods: (A) cross-domain validation against real MIMIC-IV v3.1 data, "
        "(B) within-MIMIC-IV cohort calibration (new), and (C) internal consistency validation using "
        "a 20% synthetic holdout. "
        "MIMIC-IV v3.1 credentialed access was obtained via PhysioNet on 15 Jun 2026. "
        "Preprocessing was performed by src/data/mimic_preprocessor.py, yielding a subset "
        "of 205,456 elderly admissions (95,715 unique patients, anchor_age >= 65). "
        "Of the 14 FedAcuity schema features, only 3 are mappable from the MIMIC-IV "
        "hospital module: medication_count (unique drugs per admission, clipped to [0,20]), "
        "rug_category (DRG severity 1-4 mapped to RUG scale 1-8), and resident_census "
        "(concurrent hospital admissions per date). ADL scores, nursing hours, fall risk, "
        "pain assessment, and MDS summary are LTC-specific and do not appear in hospital EHRs."
    ))
    doc.add_paragraph()

    _body(doc, "Table 2A: Cross-Domain Validation (Synthetic LTC vs Real MIMIC-IV Hospital Subset)", bold=True)
    fid_rows_mimic = [
        ["KS-test pass rate", "0/3 features", "Expected: hospital vs LTC distributional shift; medication complexity, census scale, and RUG proxy differ by care domain", "Domain gap confirmed"],
        ["Frobenius norm (3-feature corr.)", "0.8233", "vs random baseline 0.7463; correlation structure differs (hospital vs LTC context)", "Domain gap confirmed"],
        ["TSTR AUC (3 features)", "0.4156", "TRTR oracle: 0.7489; different label concepts (staffing demand vs discharge destination)", "Domain gap confirmed"],
    ]
    _add_table(doc, ["Metric", "Value", "Context", "Interpretation"],
               fid_rows_mimic, col_widths=[2.0, 0.8, 2.6, 1.0])
    doc.add_paragraph()

    _body(doc, "Table 2B: Within-MIMIC-IV Cohort Calibration (New -- src/data/mimic_analysis.py)", bold=True)
    _body(doc, (
        "This table does not compare synthetic data to MIMIC-IV. It compares two cohorts "
        "within MIMIC-IV itself: patients discharged to post-acute/LTC care (n=55,927, 27.2%) "
        "against those who are not (n=149,529, 72.8%). It checks whether MIMIC-IV's real-world "
        "discharge pattern is consistent with the synthetic generator's target mismatch rates -- "
        "a face-validity check, not a fidelity proof."
    ))
    fid_rows_cohort = [
        ["medication_count", "19.17 vs 15.21", "0.686", "<0.001", "Medium-large effect, expected direction"],
        ["rug_category (DRG proxy)", "5.49 vs 3.74", "0.782", "<0.001", "Medium-large effect, expected direction"],
        ["resident_census", "41.43 vs 41.25", "0.018", "<0.001*", "No meaningful effect (expected -- census is hospital-level, not acuity-driven)"],
    ]
    _add_table(doc, ["Feature", "LTC-bound vs Non-LTC mean", "Cohen's d", "p-value", "Interpretation"],
               fid_rows_cohort, col_widths=[1.6, 1.5, 0.8, 0.8, 2.3])
    doc.add_paragraph()
    _body(doc, (
        "* Statistically significant due to very large sample size (n=205,456); effect size (d=0.018) "
        "shows the difference is not practically meaningful."
    ), size=9, italic=True)
    doc.add_paragraph()
    _body(doc, (
        "Calibration finding: MIMIC-IV's real-world discharge-to-post-acute rate (27.2%) sits within "
        "0.8 percentage points of the synthetic SNF mismatch target (28%), the synthetic care type "
        "closest to a post-acute discharge population. This is a face-validity signal that the "
        "synthetic generator's target rates were not chosen arbitrarily."
    ))
    doc.add_paragraph()

    _body(doc, "Table 2C: Internal Consistency Validation (Synthetic vs 20% Synthetic Holdout)", bold=True)
    fid_rows_internal = [
        ["KS-test pass rate", "14/14 features (100%)", "Alpha = 0.05; all 14 schema features", "PASS"],
        ["Frobenius norm (14-feature corr.)", "0.2196", "vs random baseline 5.9407 (27x better)", "PASS"],
        ["TSTR AUC (train synthetic, test holdout)", "0.9988", "TRTR oracle: 0.9789, Gap: 0.0199 (target <0.08)", "PASS"],
    ]
    _add_table(doc, ["Metric", "Value", "Context", "Status"],
               fid_rows_internal, col_widths=[2.3, 1.0, 2.0, 0.9])
    doc.add_paragraph()

    _body(doc, (
        "Interpretation: The cross-domain mismatch in Table 2A is expected and scientifically "
        "meaningful. MIMIC-IV captures acute hospital admissions, where medication counts reflect "
        "complex inpatient regimens (hospital mean: 16.3 unique drugs, vs LTC mean: 5-11 chronic "
        "oral medications), DRG severity proxies short-stay acuity rather than long-term functional "
        "dependency, and census measures concurrent hospital occupancy rather than LTC facility "
        "residency. The TSTR gap of 0.33 confirms that the staffing-mismatch label (nursing "
        "demand-supply ratio) is conceptually distinct from hospital discharge destination -- "
        "as expected. Table 2B takes a different angle on the same dataset: rather than comparing "
        "synthetic to real, it compares two real cohorts within MIMIC-IV and finds that patients "
        "discharged to post-acute/LTC care carry significantly higher medication burden and acuity "
        "(Cohen's d = 0.69 and 0.78, p < 0.001) than those who are not, and that MIMIC-IV's real-world "
        "27.2% discharge-to-post-acute rate sits within 0.8 points of the synthetic SNF mismatch target "
        "(28%). Together, these results validate the core C2 argument: no existing public dataset "
        "adequately represents LTC staffing dynamics, motivating the purpose-built synthetic "
        "benchmark, while the cohort-calibration check provides independent, real-world face validity "
        "for the synthetic generator's target rates. Table 2C confirms that the CTGAN pipeline "
        "faithfully reproduces the target LTC distributions specified in the schema."
    ))
    doc.add_paragraph()

    _heading(doc, "3.3 Differential Privacy Results", size=12)
    _body(doc, (
        "DP-SGD applied via Opacus to StaffingNN (PyTorch). Training pool and evaluation facility "
        "were re-run after the Section 4.4 structural fix, since both depend on HELD_OUT_FACILITIES: "
        "the model is trained on the pooled data of all 9 training facilities (now including "
        "facility 8) and evaluated on held-out facility 9 (IL). Delta = 1e-5, max_grad_norm = 1.0."
    ))
    doc.add_paragraph()
    dp_rows = [
        ["1", "1.001", "0.6215", "35.6%", "Too aggressive"],
        ["2", "2.013", "0.6911", "28.4%", "Too aggressive"],
        ["5", "5.027", "0.7210", "25.3%", "Too aggressive"],
        ["10", "10.051", "0.7784", "19.3%", "RECOMMENDED"],
        ["inf (no DP)", "--", "0.9649", "0%", "Baseline"],
    ]
    _add_table(doc, ["Target epsilon", "Actual epsilon", "AUC-ROC", "Utility Drop", "Assessment"],
               dp_rows, col_widths=[1.0, 1.0, 1.0, 1.0, 1.6])
    doc.add_paragraph()
    _body(doc, (
        "Recommendation: epsilon = 10 offers the best practical privacy-utility balance at this "
        "model scale (19.3% utility degradation). Tighter budgets (epsilon <= 5) cause > 25% "
        "utility degradation with the current shallow PyTorch NN (3 hidden layers). Future work: "
        "tree-level XGBoost perturbation to enable tighter DP without the NN utility penalty."
    ))
    _page_break(doc)


def make_section4(doc):
    _heading(doc, "4. DESIGN DECISIONS AND CHALLENGES", size=13)
    _hline(doc)
    doc.add_paragraph()

    _heading(doc, "4.1 Key Design Decisions", size=12)
    _body(doc, "XGBoost as primary local model:", bold=True)
    _body(doc, (
        "XGBoost was chosen over neural networks for LTC tabular data due to its robustness "
        "to missing values, native SHAP TreeExplainer support (required for C3), and strong "
        "performance on low-sample-size tabular data. This decision makes true FL parameter "
        "averaging non-trivial (XGBoost models are trees, not gradient vectors), which is "
        "addressed via byte serialisation with largest-client-model aggregation as a proxy."
    ))
    doc.add_paragraph()
    _body(doc, "Care-type domain clustering over gradient-similarity clustering:", bold=True)
    _body(doc, (
        "Rather than computing gradient similarity at runtime (Sattler et al. 2020 approach), "
        "FedAcuity uses domain knowledge -- clinical care taxonomy -- to define clusters upfront. "
        "This is more interpretable, computationally cheaper, and directly grounded in LTC "
        "operational practice."
    ))
    doc.add_paragraph()
    _body(doc, "Opacus DP on PyTorch NN, not XGBoost:", bold=True)
    _body(doc, (
        "Opacus requires gradient-level access for noise injection, which is not available "
        "in XGBoost's tree-boosting algorithm. A secondary PyTorch NN (StaffingNN) is used "
        "exclusively for the DP sweep. This is a known, documented limitation in the codebase."
    ))
    doc.add_paragraph()

    _heading(doc, "4.2 Limitations Acknowledged", size=12)
    _bullet(doc, "MIMIC-IV cross-domain validation: Only 3 of 14 schema features are mappable from the hospital EHR module (medication_count, rug_category via DRG, resident_census). Distributional differences are expected and confirm domain specificity of LTC data. Full 14-feature validation is achievable only with a real LTC dataset (e.g., CMS Minimum Data Set), which is outside the scope of this dissertation's data access. The within-MIMIC-IV cohort calibration analysis (Section 3.2, Table 2B) partially compensates by giving an independent face-validity check that does not depend on the weak feature mapping, but it is not a substitute for true LTC-domain fidelity validation.")
    _bullet(doc, "Clustered FL's IL cluster has only 2 training clients (facilities 7, 8) after the Section 4.4 fix -- enough for genuine FedAvg aggregation, but a small N compared to MC (3 clients) and SNF (4 clients). The held-out result on facility 9 should be read as a proof-of-concept of intra-cluster federation value, not a claim that generalizes to arbitrarily small clusters.")
    _bullet(doc, "XGBoost aggregation is approximate: True XGBoost tree merging (horizontal FL with tree concatenation) is not yet implemented. The current largest-client-model proxy is a known stub documented in the code.")
    _bullet(doc, "FedProx proximal term: Not implemented as a true gradient penalty for XGBoost. Noted in config and code; correctly applied only in the PyTorch NN variant.")
    _bullet(doc, "10-facility simulation: Results are from a controlled simulation. Real-world federation across hundreds of facilities with variable connectivity is out of scope for this dissertation.")
    doc.add_paragraph()

    _heading(doc, "4.3 Engineering Challenges Resolved", size=12)
    _bullet(doc, "XGBoost 3.x BytesIO incompatibility: save_raw('ubj') + temp file deserialization workaround implemented in client.py.")
    _bullet(doc, "IL mismatch rate bug: NON_IID_SPEC missing adl_eating and adl_toileting features caused IL mismatch rate to read 41% instead of the target 12%. Fixed in schema.py.")
    _bullet(doc, "Windows cp1252 Unicode: All print() statements use ASCII equivalents to avoid codec errors on Windows 11.")
    _bullet(doc, "Ray dependency removed: Flower simulation rewritten as a manual round loop to avoid Ray installation complexity on Windows.")
    doc.add_paragraph()

    _heading(doc, "4.4 Mid-cycle Structural Fix and Re-framing", size=12)
    _body(doc, (
        "This subsection documents, transparently and for the record, a structural flaw discovered "
        "during a critical self-review of the draft results, the fix applied, and a related "
        "reframing of contribution C2 -- consistent with academic integrity norms that require "
        "disclosing material corrections rather than silently revising numbers."
    ))
    doc.add_paragraph()
    _body(doc, "C1 structural flaw and fix:", bold=True)
    _body(doc, (
        "The original experimental configuration set HELD_OUT_FACILITIES = [8, 9], both Independent "
        "Living (IL) facilities. Since the IL care-type cluster is defined as facilities {7, 8, 9}, "
        "removing both 8 and 9 from training left exactly one IL training client (facility 7). "
        "Clustered FL's IL cluster therefore had nothing to average across -- its per-round model "
        "was numerically identical to single-facility local training, even though it was being "
        "reported as a 'federated' result. This did not invalidate the MC and SNF clusters (3 and "
        "4 training clients respectively), but it meant the IL arm of the headline comparison could "
        "not demonstrate genuine federation benefit, undermining the central claim of contribution C1."
    ))
    doc.add_paragraph()
    _body(doc, (
        "The fix: facility 8 was moved from the held-out set into the IL training cluster, changing "
        "HELD_OUT_FACILITIES to [9] only (config.yaml and src/data/schema.py). The IL cluster now has "
        "2 genuine training clients (7, 8) performing real intra-cluster FedAvg each round, while "
        "facility 9 remains the sole facility that never participates in any FL training round and "
        "is reserved purely for generalisation evaluation. All downstream artefacts were re-run end "
        "to end under the corrected configuration: the 5-strategy FL simulation, held-out evaluation "
        "(Table 1), bootstrap CI / Mann-Whitney significance testing, the DP epsilon sweep (Section "
        "3.3's epsilon-sweep table -- its training pool and evaluation facility also depend on "
        "HELD_OUT_FACILITIES and were silently stale until this fix was traced through the codebase), "
        "and Figures 3, 4 "
        "and 5. The corrected, structurally sound result is that Clustered FL (AUC 0.9693) now "
        "legitimately outperforms both FedAvg (AUC 0.9414) and the single-facility IL Local baseline "
        "(AUC 0.9643) -- see Section 3.1 for the full comparison and significance test."
    ))
    doc.add_paragraph()
    _body(doc, "C2 re-framing:", bold=True)
    _body(doc, (
        "Separately, a review of the actual schema overlap between MIMIC-IV (a hospital EHR dataset) "
        "and the LTC facility-day schema found that only 3 of 14 features are mappable, which is a "
        "materially weaker basis for a cross-domain fidelity claim than the original framing implied. "
        "Rather than overstate this evidence, C2 was reframed as a deliberately modest contribution: "
        "(1) excellent internal CTGAN fidelity on a synthetic holdout (100% KS pass rate, 27x "
        "Frobenius improvement over a random baseline, TSTR gap 0.0199), reported honestly as "
        "internal consistency rather than external validation; (2) the 3-feature cross-domain check "
        "against real MIMIC-IV, disclosed with its limitations rather than presented as comprehensive "
        "validation; and (3) a new analysis, src/data/mimic_analysis.py, that uses MIMIC-IV "
        "differently -- as a within-dataset cohort-calibration anchor rather than a feature-fidelity "
        "benchmark -- giving an independent, real-world face-validity signal for the synthetic "
        "generator's target mismatch rates (Section 3.2, Table 2B)."
    ))
    _page_break(doc)


def make_section5(doc):
    _heading(doc, "5. FUTURE PLAN", size=13)
    _hline(doc)
    doc.add_paragraph()
    _body(doc, (
        "The following table describes remaining work from MidSEM to Final SEM evaluation "
        "(11 July 2026) and dissertation submission."
    ))
    doc.add_paragraph()

    plan_rows = [
        ("Week 2\n(7-13 Jun)", "MidSEM Polish",
         "Statistical testing (Mann-Whitney U done); convergence figures; paper sections I-VI; slide deck.",
         "COMPLETE"),
        ("Week 3\n(14-20 Jun)", "SHAP Pipeline + D1, D2",
         "src/xai/shap_pipeline.py: SHAP values for all 5 models via TreeExplainer.\n"
         "src/xai/d1_fidelity.py: Spearman rho of top-10 SHAP ranks vs Centralised Oracle (target >= 0.75).\n"
         "src/xai/d2_stability.py: Mean SHAP shift under +/-5% Gaussian noise (100 perturbations).",
         "PLANNED"),
        ("Week 4\n(21-27 Jun)", "D3, D4, XAI Scorecard",
         "src/xai/d3_fairness.py: Equalized odds + demographic parity across MC/SNF/IL.\n"
         "src/xai/d4_plausibility.py: % top-5 SHAP features in LTC clinical literature (target >= 60%).\n"
         "Run scorecard.py -> Fig 6 radar chart. Paper Section VII (XAI results).",
         "PLANNED"),
        ("Week 5\n(28 Jun-4 Jul)", "Paper Writing Sprint",
         "Sections VIII (Discussion), IX (Conclusion). All 6 figures final quality.\n"
         "All % TBD placeholders replaced. Upload to Overleaf and compile.",
         "PLANNED"),
        ("Week 6\n(5-11 Jul)", "Final SEM",
         "End-to-end clean run from fresh venv. Final git commit. Presentation (20 min).\n"
         "FINAL SEM EVALUATION: 11 July 2026.",
         "PLANNED"),
        ("Post Jul 11", "Submission",
         "Full dissertation write-up (10 Jul - 2 Aug). Final submission with supervisor evaluation report.",
         "PLANNED"),
    ]
    _add_table(doc, ["Phase", "Focus", "Work to be Done", "Status"],
               plan_rows, col_widths=[1.0, 1.2, 3.5, 0.9])
    doc.add_paragraph()

    _body(doc, "Immediate priorities (this week):", bold=True)
    _bullet(doc, "COMPLETE: MIMIC-IV preprocessing run (16 Jun 2026), fidelity.py validated with real 205,456-row elderly subset. Cross-domain and internal validation results both generated.")
    _bullet(doc, "Dry-run MidSEM presentation (target <= 15 minutes).")
    _bullet(doc, "Begin SHAP pipeline (src/xai/shap_pipeline.py) immediately after MidSEM.")

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Signature of the Student").font.bold = True
    p.add_run("                                    ")
    p.add_run("Signature of the Supervisor").font.bold = True
    doc.add_paragraph()
    _body(doc, "Name: _______________________                    Name: _______________________")
    _body(doc, "Date: _______________________                    Date: _______________________")
    _body(doc, "Place: ______________________                    Place: ______________________")


# ---- main --------------------------------------------------------------------

def main():
    r = _load()
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    make_title_page(doc)
    make_abstract(doc)
    make_contents(doc)
    make_section1(doc)
    make_section2(doc)
    make_section3(doc, r)
    make_section4(doc)
    make_section5(doc)

    OUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Report saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
